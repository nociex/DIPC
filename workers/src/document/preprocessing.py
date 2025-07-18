"""Document preprocessing pipeline for format detection and conversion."""

import os
import tempfile
import mimetypes
from typing import Dict, Any, Optional, Tuple, List
from pathlib import Path
from enum import Enum
import structlog

import requests
from PIL import Image
import PyPDF2
import pytesseract
from docx import Document as DocxDocument
import openpyxl

logger = structlog.get_logger(__name__)


class DocumentFormat(str, Enum):
    """Supported document formats."""
    PDF = "pdf"
    IMAGE = "image"
    DOCX = "docx"
    XLSX = "xlsx"
    TEXT = "text"
    UNKNOWN = "unknown"


class ProcessedDocument:
    """Container for processed document data."""
    
    def __init__(
        self,
        format: DocumentFormat,
        text_content: str = "",
        image_paths: List[str] = None,
        metadata: Dict[str, Any] = None,
        original_filename: str = "",
        file_size: int = 0
    ):
        self.format = format
        self.text_content = text_content
        self.image_paths = image_paths or []
        self.metadata = metadata or {}
        self.original_filename = original_filename
        self.file_size = file_size
    
    def has_text_content(self) -> bool:
        """Check if document has extractable text content."""
        return bool(self.text_content.strip())
    
    def has_images(self) -> bool:
        """Check if document has image content."""
        return len(self.image_paths) > 0
    
    def get_content_summary(self) -> Dict[str, Any]:
        """Get summary of document content."""
        return {
            "format": self.format.value,
            "has_text": self.has_text_content(),
            "text_length": len(self.text_content),
            "image_count": len(self.image_paths),
            "metadata": self.metadata,
            "original_filename": self.original_filename,
            "file_size": self.file_size
        }


class DocumentPreprocessor:
    """Handles document format detection and preprocessing."""
    
    def __init__(self, temp_dir: Optional[str] = None):
        """
        Initialize preprocessor.
        
        Args:
            temp_dir: Directory for temporary files. If None, uses system temp.
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self.supported_image_formats = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
        self.supported_text_formats = {'.txt', '.md', '.csv'}
    
    def detect_format(self, file_path: str, content_type: Optional[str] = None) -> DocumentFormat:
        """
        Detect document format from file path and content type.
        
        Args:
            file_path: Path to the document file
            content_type: MIME type if available
            
        Returns:
            Detected document format
        """
        file_path_lower = file_path.lower()
        
        # Check content type first if available
        if content_type:
            if content_type.startswith('image/'):
                return DocumentFormat.IMAGE
            elif content_type == 'application/pdf':
                return DocumentFormat.PDF
            elif 'word' in content_type or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return DocumentFormat.DOCX
            elif 'sheet' in content_type or content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                return DocumentFormat.XLSX
            elif content_type.startswith('text/'):
                return DocumentFormat.TEXT
        
        # Fallback to file extension
        if file_path_lower.endswith('.pdf'):
            return DocumentFormat.PDF
        elif any(file_path_lower.endswith(ext) for ext in self.supported_image_formats):
            return DocumentFormat.IMAGE
        elif file_path_lower.endswith('.docx'):
            return DocumentFormat.DOCX
        elif file_path_lower.endswith(('.xlsx', '.xls')):
            return DocumentFormat.XLSX
        elif any(file_path_lower.endswith(ext) for ext in self.supported_text_formats):
            return DocumentFormat.TEXT
        
        return DocumentFormat.UNKNOWN
    
    def download_file(self, file_url: str) -> Tuple[str, str]:
        """
        Download file from URL to temporary location.
        
        Args:
            file_url: URL of the file to download
            
        Returns:
            Tuple of (local_file_path, content_type)
            
        Raises:
            Exception: If download fails
        """
        try:
            logger.info("Downloading file for processing", file_url=file_url)
            
            response = requests.get(file_url, timeout=300, stream=True)
            response.raise_for_status()
            
            # Get content type
            content_type = response.headers.get('content-type', 'application/octet-stream')
            
            # Generate temporary file path
            file_extension = mimetypes.guess_extension(content_type) or ''
            temp_file = tempfile.NamedTemporaryFile(
                delete=False, 
                suffix=file_extension,
                dir=self.temp_dir
            )
            
            # Download file in chunks
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    temp_file.write(chunk)
            
            temp_file.close()
            
            logger.info(
                "File downloaded successfully",
                file_url=file_url,
                local_path=temp_file.name,
                content_type=content_type,
                file_size=os.path.getsize(temp_file.name)
            )
            
            return temp_file.name, content_type
            
        except Exception as e:
            logger.error("Failed to download file", file_url=file_url, error=str(e))
            raise Exception(f"File download failed: {str(e)}")
    
    def process_pdf(self, file_path: str) -> ProcessedDocument:
        """
        Process PDF document to extract text and convert pages to images.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ProcessedDocument with extracted content
        """
        try:
            logger.info("Processing PDF document", file_path=file_path)
            
            text_content = ""
            image_paths = []
            metadata = {}
            
            # Extract text using PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}", error=str(e))
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    })
            
            # Convert PDF pages to images for vision models
            try:
                from pdf2image import convert_from_path
                
                # Convert first few pages to images (limit to avoid excessive processing)
                max_pages = 5
                images = convert_from_path(
                    file_path, 
                    first_page=1, 
                    last_page=min(max_pages, metadata.get('page_count', 1))
                )
                
                for i, image in enumerate(images):
                    image_path = os.path.join(
                        self.temp_dir, 
                        f"pdf_page_{i+1}_{os.path.basename(file_path)}.png"
                    )
                    image.save(image_path, 'PNG')
                    image_paths.append(image_path)
                    
                logger.info(f"Converted {len(images)} PDF pages to images")
                
            except ImportError:
                logger.warning("pdf2image not available, skipping PDF to image conversion")
            except Exception as e:
                logger.warning("Failed to convert PDF to images", error=str(e))
            
            return ProcessedDocument(
                format=DocumentFormat.PDF,
                text_content=text_content.strip(),
                image_paths=image_paths,
                metadata=metadata,
                original_filename=os.path.basename(file_path),
                file_size=os.path.getsize(file_path)
            )
            
        except Exception as e:
            logger.error("Failed to process PDF", file_path=file_path, error=str(e))
            raise Exception(f"PDF processing failed: {str(e)}")
    
    def process_image(self, file_path: str) -> ProcessedDocument:
        """
        Process image document, optionally extracting text via OCR.
        
        Args:
            file_path: Path to image file
            
        Returns:
            ProcessedDocument with image and optional OCR text
        """
        try:
            logger.info("Processing image document", file_path=file_path)
            
            text_content = ""
            metadata = {}
            
            # Load image and get metadata
            with Image.open(file_path) as img:
                metadata.update({
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode
                })
                
                # Attempt OCR text extraction
                try:
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        text_content = ocr_text.strip()
                        metadata['ocr_extracted'] = True
                        logger.info("OCR text extraction successful", text_length=len(text_content))
                    else:
                        metadata['ocr_extracted'] = False
                        logger.info("No text found via OCR")
                        
                except Exception as e:
                    logger.warning("OCR text extraction failed", error=str(e))
                    metadata['ocr_extracted'] = False
                    metadata['ocr_error'] = str(e)
            
            return ProcessedDocument(
                format=DocumentFormat.IMAGE,
                text_content=text_content,
                image_paths=[file_path],
                metadata=metadata,
                original_filename=os.path.basename(file_path),
                file_size=os.path.getsize(file_path)
            )
            
        except Exception as e:
            logger.error("Failed to process image", file_path=file_path, error=str(e))
            raise Exception(f"Image processing failed: {str(e)}")
    
    def process_docx(self, file_path: str) -> ProcessedDocument:
        """
        Process DOCX document to extract text content.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ProcessedDocument with extracted text
        """
        try:
            logger.info("Processing DOCX document", file_path=file_path)
            
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text += " | ".join(row_text) + "\n"
            
            if table_text:
                text_content += "\n--- Tables ---\n" + table_text
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'has_tables': len(doc.tables) > 0
            }
            
            return ProcessedDocument(
                format=DocumentFormat.DOCX,
                text_content=text_content.strip(),
                metadata=metadata,
                original_filename=os.path.basename(file_path),
                file_size=os.path.getsize(file_path)
            )
            
        except Exception as e:
            logger.error("Failed to process DOCX", file_path=file_path, error=str(e))
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    def process_xlsx(self, file_path: str) -> ProcessedDocument:
        """
        Process XLSX document to extract text content from cells.
        
        Args:
            file_path: Path to XLSX file
            
        Returns:
            ProcessedDocument with extracted text
        """
        try:
            logger.info("Processing XLSX document", file_path=file_path)
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            text_content = ""
            sheet_data = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = f"\n--- Sheet: {sheet_name} ---\n"
                
                # Extract data from cells
                rows_with_data = 0
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None and str(cell).strip() for cell in row):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell).strip())
                            else:
                                row_text.append("")
                        sheet_text += " | ".join(row_text) + "\n"
                        rows_with_data += 1
                
                if rows_with_data > 0:
                    text_content += sheet_text
                    sheet_data[sheet_name] = rows_with_data
            
            metadata = {
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'sheet_data': sheet_data
            }
            
            return ProcessedDocument(
                format=DocumentFormat.XLSX,
                text_content=text_content.strip(),
                metadata=metadata,
                original_filename=os.path.basename(file_path),
                file_size=os.path.getsize(file_path)
            )
            
        except Exception as e:
            logger.error("Failed to process XLSX", file_path=file_path, error=str(e))
            raise Exception(f"XLSX processing failed: {str(e)}")
    
    def process_text(self, file_path: str) -> ProcessedDocument:
        """
        Process plain text document.
        
        Args:
            file_path: Path to text file
            
        Returns:
            ProcessedDocument with text content
        """
        try:
            logger.info("Processing text document", file_path=file_path)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = ""
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                        encoding_used = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                raise Exception("Could not decode text file with any supported encoding")
            
            metadata = {
                'encoding': encoding_used,
                'line_count': len(text_content.splitlines()),
                'character_count': len(text_content)
            }
            
            return ProcessedDocument(
                format=DocumentFormat.TEXT,
                text_content=text_content,
                metadata=metadata,
                original_filename=os.path.basename(file_path),
                file_size=os.path.getsize(file_path)
            )
            
        except Exception as e:
            logger.error("Failed to process text file", file_path=file_path, error=str(e))
            raise Exception(f"Text processing failed: {str(e)}")
    
    def process_document(self, file_url: str) -> ProcessedDocument:
        """
        Main method to process a document from URL.
        
        Args:
            file_url: URL of the document to process
            
        Returns:
            ProcessedDocument with extracted content
            
        Raises:
            Exception: If processing fails
        """
        local_file_path = None
        
        try:
            # Download file
            local_file_path, content_type = self.download_file(file_url)
            
            # Detect format
            document_format = self.detect_format(local_file_path, content_type)
            
            logger.info(
                "Processing document",
                file_url=file_url,
                local_path=local_file_path,
                detected_format=document_format.value,
                content_type=content_type
            )
            
            # Process based on format
            if document_format == DocumentFormat.PDF:
                return self.process_pdf(local_file_path)
            elif document_format == DocumentFormat.IMAGE:
                return self.process_image(local_file_path)
            elif document_format == DocumentFormat.DOCX:
                return self.process_docx(local_file_path)
            elif document_format == DocumentFormat.XLSX:
                return self.process_xlsx(local_file_path)
            elif document_format == DocumentFormat.TEXT:
                return self.process_text(local_file_path)
            else:
                # Try to process as text first, then as image
                try:
                    return self.process_text(local_file_path)
                except:
                    try:
                        return self.process_image(local_file_path)
                    except:
                        raise Exception(f"Unsupported document format: {document_format.value}")
        
        finally:
            # Clean up downloaded file
            if local_file_path and os.path.exists(local_file_path):
                try:
                    os.unlink(local_file_path)
                    logger.debug("Cleaned up temporary file", file_path=local_file_path)
                except Exception as e:
                    logger.warning("Failed to clean up temporary file", file_path=local_file_path, error=str(e))
    
    def cleanup_temp_files(self, file_paths: List[str]):
        """
        Clean up temporary files.
        
        Args:
            file_paths: List of file paths to clean up
        """
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.unlink(file_path)
                    logger.debug("Cleaned up temporary file", file_path=file_path)
            except Exception as e:
                logger.warning("Failed to clean up temporary file", file_path=file_path, error=str(e))